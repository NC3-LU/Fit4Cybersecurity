import os
import logging
from django.http import HttpResponse
from django.conf import settings

from csskp.settings import PICTURE_DIR
from survey.models import (
    SurveyQuestion,
    SurveyQuestionAnswer,
    SurveyUser,
    SurveyUserAnswer,
    Recommendations,
    TranslationKey,
)
from survey.globals import TRANSLATION_UI
from utils.radarFactory import radar_factory
import matplotlib.pyplot as plt

# remove html tags from the report https://stackoverflow.com/questions/753052/strip-html-from-strings-in-python
from io import StringIO
from html.parser import HTMLParser

import docx
from docx import Document
from docx.shared import Cm, Pt

# from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import re

# Get an instance of a logger
logger = logging.getLogger(__name__)


class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = StringIO()

    def handle_data(self, d):
        self.text.write(d)

    def get_data(self):
        return self.text.getvalue()


def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()


def getRecommendations(user: SurveyUser, lang: str):
    allAnswers = SurveyQuestionAnswer.objects.all().order_by(
        "question__qindex", "aindex"
    )
    recommendations_translations = get_formatted_translations(lang, "R")
    categories_translations = get_formatted_translations(lang, "C")

    finalReportRecs = {}

    for a in allAnswers:
        userAnswer = SurveyUserAnswer.objects.filter(user=user).filter(answer=a)[0]
        recommendations = Recommendations.objects.filter(forAnswer=a)

        if not recommendations.exists():
            continue

        for rec in recommendations:
            if rec.min_e_count > user.e_count or rec.max_e_count < user.e_count:
                continue
            if (userAnswer.uvalue > 0 and rec.answerChosen) or (
                userAnswer.uvalue <= 0 and not rec.answerChosen
            ):
                category_name = categories_translations[
                    rec.forAnswer.question.service_category.titleKey
                ]

                translated_recommendation = recommendations_translations[rec.textKey]
                if is_recommendation_already_added(
                    translated_recommendation, finalReportRecs
                ):
                    continue

                if category_name not in finalReportRecs:
                    finalReportRecs[category_name] = []
                finalReportRecs[category_name].append(
                    recommendations_translations[rec.textKey]
                )

    return finalReportRecs


def is_recommendation_already_added(recommendation: str, recommendations: dict):
    if recommendations:
        for category, recommendations_per_category in recommendations.items():
            if recommendation in recommendations_per_category:
                return True

    return False


def createAndSendReport(user: SurveyUser, lang: str):
    """Generates the report as a .docx file, then returns it to the view."""

    filepath = settings.BASE_DIR + "/wtemps/"

    template = filepath + "template-" + lang + ".docx"
    doc = Document(template)

    introduction = ""
    file_path = os.path.join(filepath, lang + "_intro.txt")
    try:
        with open(file_path, encoding="UTF-8") as f:
            introduction = f.read()
    except Exception as e:
        logger.error("Error when reading template: {}".format(e))
        raise Exception("Oooops... something went wrong!")

    introduction = introduction.replace("\n\r", "\n")
    introduction = introduction.split("\n\n")
    x = 0
    for i in introduction:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1
            continue
        doc.add_paragraph(i)

    methodDescr = ""
    file_path = os.path.join(filepath, lang + "_description.txt")
    try:
        with open(file_path, encoding="UTF-8") as f:
            methodDescr = f.read()
    except Exception as e:
        logger.error("Error when reading template: {}".format(e))
        raise Exception("Oooops... something went wrong!")

    methodDescr = methodDescr.replace("\n\r", "\n")
    methodDescr = methodDescr.replace(
        "$$qnumber$$", str(len(SurveyQuestion.objects.all()))
    )
    methodDescr = methodDescr.split("\n\n")
    x = 0
    for i in methodDescr:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1
            continue
        doc.add_paragraph(i)

    results = ""
    file_path = os.path.join(filepath, lang + "_resultdisclaimer.txt")
    try:
        with open(file_path, encoding="UTF-8") as f:
            results = f.read()
    except Exception as e:
        logger.error("Error when reading template: {}".format(e))
        raise Exception("Oooops... something went wrong!")

    score, details, section_list = calculateResult(user, lang)

    results = results.replace("\n\r", "\n")
    # results = results.replace("$$result$$", str(score))
    results = results.split("\n\n")

    x = 0
    for i in results:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1

            continue
        if "$$result$$" in i:
            i = i.split("$$result$$")
            p = doc.add_paragraph()
            ind = 1
            for x in i:
                p.add_run(x)
                # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if ind < len(i):
                    px = p.add_run(str(score))
                    px.font.bold = True
                ind += 1
        else:
            doc.add_paragraph(i)

    try:
        chart_png_file = generate_chart_png(user, details, section_list, lang)
        doc.add_paragraph()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(chart_png_file)
    except:
        pass

    doc.add_paragraph()

    recommendationList = getRecommendations(user, lang)

    doc.add_page_break()

    recs = ""
    file_path = os.path.join(filepath, lang + "_recs.txt")
    try:
        with open(file_path, encoding="UTF-8") as f:
            recs = f.read()
    except Exception as e:
        logger.error("Error when reading template: {}".format(e))
        raise Exception("Oooops... something went wrong!")

    recs = recs.replace("\n\r", "\n")
    recs = recs.split("\n\n")
    x = 0
    for i in recs:
        if x == 0:
            doc.add_heading(i, level=1)
            x += 1
            continue
        doc.add_paragraph(i)

    for category, items in recommendationList.items():
        doc.add_heading(category, level=2)
        point_number = 1
        for recommendation in items:
            # Create hyperlinks in the document.
            split_links = re.split(
                '\<a href="(.+?)".*>([\w\s]+)\<\/a\>', recommendation
            )
            elements_number = len(split_links)
            if elements_number > 1:
                paragraph = doc.add_paragraph(
                    str(point_number) + ". ", "List Paragraph"
                )
                for index, string_part in enumerate(split_links):
                    if index % 3 == 0:
                        paragraph.add_run(string_part)
                        if elements_number > index + 1:
                            add_hyperlink(
                                paragraph,
                                split_links[index + 2],
                                split_links[index + 1],
                            )
            else:
                doc.add_paragraph(
                    str(point_number) + ". " + strip_tags(recommendation),
                    "List Paragraph",
                )
            point_number += 1

    doc.add_page_break()

    doc.add_heading(TRANSLATION_UI["document"]["questions"][lang], level=1)

    questions_translations = get_formatted_translations(lang, "Q")
    answers_translations = get_formatted_translations(lang, "A")

    questions = SurveyQuestion.objects.all().order_by("qindex")

    for index, question in enumerate(questions):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = str(index + 1)
        hdr_cells[1].text = strip_tags(questions_translations[question.titleKey])

        bX = hdr_cells[0].paragraphs[0].runs[0]
        bX.font.bold = True
        bX.font.size = Pt(13)
        bX = hdr_cells[1].paragraphs[0].runs[0]
        bX.font.bold = True
        bX.font.size = Pt(13)

        answers_list = SurveyQuestionAnswer.objects.filter(question=question).order_by(
            "aindex"
        )
        user_answers = SurveyUserAnswer.objects.filter(user=user)
        user_answers_values = {}
        for user_answer in user_answers:
            user_answers_values[user_answer.answer_id] = user_answer

        for answer in answers_list:
            row_cells = table.add_row().cells

            user_answer = user_answers_values[answer.id]
            if user_answer.uvalue > 0:
                row_cells[0].text = "X"
                bX = row_cells[0].paragraphs[0].runs[0]
                bX.font.bold = True
            else:
                row_cells[0].text = " "

            row_cells[1].text = strip_tags(answers_translations[answer.answerKey])

            if user_answer.uvalue > 0:
                bX = row_cells[1].paragraphs[0].runs[0]
                bX.font.bold = True
                if user_answer.content != "":
                    content_row_cells = table.add_row().cells
                    content_row_cells[0].text = " "
                    content_row_cells[1].text = strip_tags(user_answer.content)

        col = table.columns[0]
        col.width = Cm(1.0)
        col = table.columns[1]
        col.width = Cm(14.0)
        for cell in table.columns[0].cells:
            cell.width = Cm(1.0)
        for cell in table.columns[1].cells:
            cell.width = Cm(14.0)

        doc.add_paragraph()

    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = str(date.today()) + "\t\tFit4Cybersecurity"
    paragraph.style = doc.styles["Header"]

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        "attachment; filename=Report_Fit4Cybersecurity_"
        + str(date.today())
        + "_"
        + lang
        + ".docx"
    )
    doc.save(response)

    # make checkboxes to recommendation and a single button of get companies
    # then call getcompanies when button is hit

    return response


def calculateResult(user: SurveyUser, lang: str):
    total_questions_score = 0
    total_user_score = 0
    user_evaluations_per_section = {}
    max_evaluations_per_section = {}
    sections_list = []

    translation_key_values = get_formatted_translations(lang, "S")

    for question in SurveyQuestion.objects.all():
        total_questions_score += question.maxPoints

        if question.section.id not in max_evaluations_per_section:
            max_evaluations_per_section[question.section.id] = 0
        max_evaluations_per_section[question.section.id] += question.maxPoints

        section_title = translation_key_values[question.section.sectionTitleKey]
        if section_title not in sections_list:
            sections_list.append(section_title)

    user_selected_answers = SurveyUserAnswer.objects.filter(
        user=user, uvalue__gt=0
    ).order_by("answer__question__qindex", "answer__aindex")
    for user_selected_answer in user_selected_answers:
        section_id = user_selected_answer.answer.question.section.id

        total_user_score += user_selected_answer.answer.score

        if section_id not in user_evaluations_per_section:
            user_evaluations_per_section[section_id] = 0
        user_evaluations_per_section[section_id] += user_selected_answer.answer.score

    # get the score in percent! with then 100 being total_questions_score
    total_user_score = round(total_user_score * 100 / total_questions_score)

    user_evaluations = []
    for section_id, user_evaluation_per_section in user_evaluations_per_section.items():
        user_evaluations.append(
            round(
                user_evaluation_per_section
                * 100
                / max_evaluations_per_section[section_id]
            )
        )

    return total_user_score, user_evaluations, sections_list


def generate_chart_png(user: SurveyUser, evaluation, sections_list, lang):
    """Generates the chart with Matplotlib and returns the path of the generated
    graph which will be included in the report.
    """
    n = len(sections_list)
    theta = radar_factory(n, frame="polygon")

    spoke_labels = []
    for section in sections_list:
        spoke_labels.append(section)

    fig, ax = plt.subplots(
        figsize=(7, 5), dpi=150, nrows=1, ncols=1, subplot_kw=dict(projection="radar")
    )
    fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)

    ax.set_rgrids([0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100])
    ax.set_ylim(0, 100)

    ax.plot(theta, evaluation, color="r")
    ax.fill(theta, evaluation, facecolor="r", alpha=0.25)

    ax.set_varlabels(spoke_labels)

    ax.legend(
        [TRANSLATION_UI["report"]["result"][lang]],
        loc=(0.9, 0.95),
        labelspacing=0.1,
        fontsize="small",
    )

    fig.text(
        1.0,
        1.0,
        TRANSLATION_UI["report"]["chart"][lang],
        horizontalalignment="center",
        color="black",
        weight="bold",
        size="large",
    )

    if not os.path.isdir(PICTURE_DIR):
        os.makedirs(PICTURE_DIR)
    file_name = os.path.join(PICTURE_DIR, "survey-{}.png".format(user.user_id))
    try:
        res = plt.savefig(file_name)
    except Exception as e:
        return ""
    finally:
        plt.close()

    return file_name


def get_formatted_translations(lang: str, type: str):
    translations = TranslationKey.objects.filter(lang=lang, ttype=type)
    translation_key_values = {}
    for translation in translations:
        translation_key_values[translation.key] = translation.text

    return translation_key_values


def add_hyperlink(paragraph, text: str, url: str):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Set link color and underline
    c = docx.oxml.shared.OxmlElement("w:color")
    c.set(docx.oxml.shared.qn("w:val"), "FF8822")
    rPr.append(c)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
