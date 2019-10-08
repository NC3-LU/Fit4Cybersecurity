from django.http import HttpResponse
from django.conf import settings

from survey.models import SurveyQuestion, SurveyQuestionAnswer, SurveyUser, SurveyUserAnswer, Recommendations, \
    TranslationKey
from survey.globals import SECTOR_CHOICES, COMPANY_SIZE, TRANSLATION_UI
from utils.radarFactory import radar_factory
import matplotlib.pyplot as plt


def getRecommendations(cuser):
    allAnswers = SurveyQuestionAnswer.objects.all().order_by('question__qindex','aindex')

    #userAnswers = SurveyUserAnswer.objects.all().filter(user=cuser)

    #recommendations = Recommendations

    finalReportRecs = []

    for a in allAnswers:
        userAnswer = SurveyUserAnswer.objects.all().filter(user=cuser).filter(answer=a)[0]
        recommendation = Recommendations.objects.all().filter(forAnswer=a)

        if not recommendation.exists():
            continue

        for rec in recommendation:
            if rec.min_e_count.lower() > cuser.e_count.lower() or rec.max_e_count.lower() < cuser.e_count.lower():
                continue
            if userAnswer.uvalue > 0 and rec.answerChosen:
                finalReportRecs.append(str(rec))
            elif userAnswer.uvalue <= 0 and not rec.answerChosen:
                finalReportRecs.append(str(rec))

    return finalReportRecs


def createAndSendReport(userID, lang):
    from datetime import date
    from docx import Document
    from docx.shared import Cm, Inches, Pt

    cuser = SurveyUser.objects.filter(user_id=userID)[0]

    filepath = settings.BASE_DIR+"/wtemps/"

    theImage = filepath+"monarc.jpg"
    template = filepath+lang.lower()+"1.docx"
    doc = Document(template)
    #document = MailMerge(template)
    #doc = Document()

    score = 80
    score, detailMax, details, section_list = calculateResult(cuser)

    everyQuestion = SurveyQuestion.objects.all().order_by('qindex')

    sectorName = str(cuser.sector)
    for a,b in SECTOR_CHOICES:
        if cuser.sector == a:
            sectorName = str(b)
    
    compSize = str(cuser.e_count)
    for a,b in COMPANY_SIZE:
        if cuser.e_count == a:
            compSize = b
    
    recommendationList = getRecommendations(cuser)
    recommendationList = "\n\n".join(recommendationList)

    tfile = open(filepath+"/"+lang.lower()+"_intro.txt",'r')
    introduction = tfile.read()
    tfile.close()

    introduction = introduction.replace("\n\r","\n")
    introduction = introduction.split("\n\n")
    x = 0
    for i in introduction:
        if x == 0:
            doc.add_heading(i,level=1)
            x += 1
            continue
        doc.add_paragraph(i)


    tfile = open(filepath+"/"+lang.lower()+"_description.txt",'r')
    methodDescr = tfile.read()
    tfile.close()
    
    methodDescr = methodDescr.replace("\n\r","\n")
    methodDescr = methodDescr.split("\n\n")
    x = 0
    for i in methodDescr:
        if x == 0:
            doc.add_heading(i,level=1)
            x += 1
            continue
        doc.add_paragraph(i)
    

    tfile = open(filepath+"/"+lang.lower()+"_resultdisclaimer.txt",'r')
    results = tfile.read()
    tfile.close()

    results = results.replace("\n\r","\n")
    results = results.replace("$$result$$",str(score))
    results = results.split("\n\n")

    x = 0
    for i in results:
        if x == 0:
            doc.add_heading(i,level=1)
            x += 1
            continue
        doc.add_paragraph(i)


    doc.add_heading(TRANSLATION_UI['document']['questions'][lang.lower()],level=1)

    x = 1
    for i in everyQuestion:
        table = doc.add_table(rows=1,cols=2)
        table.autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = str(x)
        hdr_cells[1].text = str(i)

        bX = hdr_cells[0].paragraphs[0].runs[0]
        bX.font.bold = True
        bX.font.size = Pt(13)
        bX = hdr_cells[1].paragraphs[0].runs[0]
        bX.font.bold = True
        bX.font.size = Pt(13)

        answerlist = SurveyQuestionAnswer.objects.filter(question=i).order_by('aindex')
        
        for a in answerlist:
            row_cells = table.add_row().cells
            u = SurveyUserAnswer.objects.filter(answer=a)[0]
            
            if u.uvalue > 0:
                row_cells[0].text = "X"
                bX = row_cells[0].paragraphs[0].runs[0]
                bX.font.bold = True
            else:
                row_cells[0].text = " "
            
            row_cells[1].text = str(a)
        
        col = table.columns[0]
        col.width = Cm(1.5)
        col = table.columns[1]
        col.width = Cm(14.0)
        for cell in table.columns[0].cells:
            cell.width=Cm(1.5)

        doc.add_paragraph()
        x += 1

    # todo: add chart_png_file = generate_chart_png(user, detailMax, details, section_list) to the file.


    '''
    table = []
    ind = 0
    for i in everyQuestion:
        ind += 1
        if ind > 1:
            table.append({'ca':"", 'surveyAnswers':""})

        answerlist = SurveyQuestionAnswer.objects.filter(question=i).order_by('aindex')
        headingLine = {'ca':str(ind), 'surveyAnswers':str(i)}
        table.append(headingLine)
        
        for a in answerlist:
            line = {'ca':"", 'surveyAnswers':""}
            u = SurveyUserAnswer.objects.filter(answer=a)[0]
            
            if u.uvalue > 0:
                line['ca'] = "X"
            else:
                line['ca'] = " "
            
            line['surveyAnswers'] = str(a)
            table.append(line)

    everyQuestionAndAnswer = table
    '''
    '''
    document.merge(
        result=str(theResult)+"/100",
        companysize=compSize,
        resultGraph=theImage,
        #surveyAnswers=everyQuestionAndAnswer,
        ca=everyQuestionAndAnswer,
        sector=sectorName,
        generationDate=str(date.today()),
        recommendationsList=recommendationList,
        )
    '''

    # use matplotlib for png of radar graph
    # can use matplotlib import pyplot as plt
    # then the graph save: plt.savefig('/tmp/'+str(userID)+'.png')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=result-'+lang.lower()+'.docx'
    #document.write(response)
    doc.save(response)
    # make survey readonly and show results.
    # make checkboxes to recommendation and a single button of get companies
    # then call getcompanies when button is hit

    return response


def calculateResult(user: SurveyUser):
    allUserAnswers = SurveyUserAnswer.objects.filter(uvalue__gt=0, user=user).order_by('answer__question__qindex',
                                                                                       'answer__aindex')
    totalscore = sum([x.answer.score for x in allUserAnswers])

    maxeval = {}
    evaluation = {}
    sectionlist = {}
    maxscore = 0

    translations = TranslationKey.objects.filter(lang=user.chosenLang, ttype='S')
    translation_key_values = {}
    for translation in translations:
        translation_key_values[translation.key] = translation.text

    for q in SurveyQuestion.objects.all():
        maxscore += q.maxPoints
        if q.section.id not in evaluation:
            evaluation[q.section.id] = 0
        if q.section.id not in maxeval:
            maxeval[q.section.id] = 0
        sectionlist[q.section.id] = translation_key_values[q.section.sectionTitleKey]

        maxeval[q.section.id] += q.maxPoints

        uanswers = SurveyUserAnswer.objects.filter(user=user, uvalue__gt=0, answer__question__id=q.id)
        scores = [x.answer.score for x in uanswers]
        evaluation[q.section.id] += sum(scores)

    # get the score in percent! with then 100 being maxscore
    totalscore = round((totalscore * 100) / maxscore)

    sectionlist = [sectionlist[x] for x in sectionlist]
    evaluation = [evaluation[x] for x in evaluation]
    maxeval = [maxeval[x] for x in maxeval]

    return totalscore, maxeval, evaluation, sectionlist


def generate_chart_png(user: SurveyUser, max_eval, evaluation, sections_list):
    n = len(sections_list)
    theta = radar_factory(n, frame='polygon')

    spoke_labels = []
    for section in sections_list:
        spoke_labels.append(section)

    fig, ax = plt.subplots(figsize=(n, n), nrows=1, ncols=1,
                             subplot_kw=dict(projection='radar'))
    fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)

    colors = ['b', 'r', 'g', 'm', 'y', 'c', 'k', 'w']

    # data = [
    #     ('Basecase', [
    #         [0.88, 0.01],
    #         [0.07, 0.95],
    #     ]),
    # ]
    # for ax, (title, case_data) in zip(axes.flat, data):
    #     ax.set_rgrids([0.2, 0.4, 0.6, 0.8])
    #     ax.set_title(title, weight='bold', size='medium', position=(0.5, 1.1),
    #                  horizontalalignment='center', verticalalignment='center')
    #     for d, color in zip(case_data, colors):
    #         ax.plot(theta, d, color=color)
    #         ax.fill(theta, d, facecolor=color, alpha=0.25)
    #     ax.set_varlabels(spoke_labels)


    grid_step = max_eval[1] / 5
    ax.set_rgrids([grid_step, grid_step * 2, grid_step * 3, grid_step * 4])

    ax.plot(theta, evaluation, color=colors[1])
    ax.fill(theta, evaluation, facecolor=colors[1], alpha=0.25)
    ax.set_varlabels(spoke_labels)

    # add legend relative to top-left plot
    ax.legend(sections_list, loc=(0.9, .95),
                       labelspacing=0.1, fontsize='small')

    fig.text(0.5, 0.965, TRANSLATION_UI['report']['chart'][user.chosenLang.lower()],
             horizontalalignment='center', color='black', weight='bold',
             size='large')

    file_name = './static/users/survey-' + str(user.id) + '.png'
    plt.savefig(file_name)

    return file_name
